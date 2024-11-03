
# Create your views here.

from django.shortcuts import redirect, render , get_object_or_404
from .forms import OrderForm,  CustomerForm , BillingForm , SupplierForm , LorryForm , ProductForm
from .models import Orders, Customer, Product, TempRate , Billing ,  Record , TempTable , OrderGroupReference , Supplier , Lorry , Payment
from django.db.models import Q, Case, When, IntegerField
import csv
from django.http import HttpResponse, JsonResponse
from django.views.decorators.http import require_GET
from .forms import calculate_amount
from datetime import datetime , timedelta
from django.db.models import Sum, F, Avg
from django.utils.datetime_safe import datetime
from django.utils import timezone
from django.views.decorators.http import require_POST
from django.db.models import  Min ,Max
from django.db.models.functions import Round
import itertools
from django.db import transaction
import datetime as inv_datetime
from django.db.models import Count
from mailmerge import MailMerge
from django.template.loader import get_template
from django.http import FileResponse
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os
from num2words import num2words
from django.db.models import Value, CharField
from datetime import datetime
from django.db import transaction
import logging
from django.conf import settings
from django.utils.dateparse import parse_date
from itertools import groupby
from operator import itemgetter
from django.template import TemplateDoesNotExist
from collections import defaultdict


def orderDocView(request):
    return render(request, 'invoice/orderDoc.html')

def rateDocView(request):
    return render(request, 'invoice/rateDoc.html')
def welcomeView(request):
    return render(request, 'invoice/welcome.html')

def get_lorry_details(request, lorry_id):
    try:
        lorry = Lorry.objects.get(Lno=lorry_id)
        data = {
            'height': lorry.height,
            'width': lorry.width,
            'length': lorry.length,
        }
        return JsonResponse(data)
    except Lorry.DoesNotExist:
        return JsonResponse({'error': 'Lorry not found'}, status=404)

def orderFormView(request):
    form = OrderForm()
    last_record = Record.objects.order_by('record_id').last()
    if last_record and last_record.Invoice:
        # Extract the part of the Invoice after the first slash
        invoice_parts = last_record.Invoice.split('/')
        if len(invoice_parts) > 1:
            # Increment the number after the first slash
            new_invoice_number = str(int(invoice_parts[1]) + 1)
        else:
            new_invoice_number = '1'  # Default to 1 if no valid number exists
    else:
        new_invoice_number = '1'



    if request.method == 'POST':
        form = OrderForm(request.POST)


        if form.is_valid():
            order_instance = form.save(commit=False)

            if not order_instance.df2:
                order_instance.df2 = order_instance.df1
            if not order_instance.df3:
                order_instance.df3 = order_instance.df1
            # Call the calculate_amount function and set the amount field
            order_instance.amount = calculate_amount(order_instance)

            # Save the order instance with the calculated amount
            order_instance.save()

            return redirect('show_url')
        else:
            # If the form is not valid, you can log the errors
            print(form.errors)

    template_name = 'invoice/order.html'
    context = {'form': form , 'new_invoice_number': new_invoice_number}
    return render(request, template_name, context)



def showView(request):
    supplier_query = request.GET.get('q', '')
    customer_query = request.GET.get('customer', '')
    product_query = request.GET.get('product', '')
    start_date_query = request.GET.get('start_date', '')
    end_date_query = request.GET.get('end_date', '')

    suppliers = Supplier.objects.all()
    customers = Customer.objects.all()
    products = Product.objects.all()

    total_obj = Orders.objects.all()
    total_count = total_obj.count()

    obj = Orders.objects.all()

    if supplier_query:
        obj = obj.filter(Q(fname__fname__icontains=supplier_query))
    if customer_query:
        obj = obj.filter(Q(Cname__Cname__icontains=customer_query))
    if product_query:
        obj = obj.filter(Q(product__name__icontains=product_query))
    if start_date_query and end_date_query:
        start_date = parse_date(start_date_query)
        end_date = parse_date(end_date_query)
        if start_date and end_date:
            obj = obj.filter(df3__range=[start_date, end_date])
    obj = obj.order_by('df3')

    filtered_count = obj.count()

    template_name = 'invoice/show.html'
    context = { 'obj': obj,
                'query': supplier_query,
                'customer_query': customer_query,
                'product_query': product_query,
                'start_date_query': start_date_query,
                'end_date_query': end_date_query,
                'filtered_count': filtered_count,
                'total_count': total_count,
                'suppliers': suppliers,
                'customers': customers,
                'products': products,}
    return render(request, template_name, context)

def unpaidView(request, order_id):
    # Only proceed if the request is a POST request
    if request.method == "POST":
        # Get the order object based on the order_id (using oid)
        order = get_object_or_404(Orders, oid=order_id)

        # Calculate the payment date as the 1st day of the next month of df3
        if order.df3:
            year = order.df3.year
            month = order.df3.month
            # If it's December, set to January of the next year
            if month == 12:
                payment_date = timezone.datetime(year + 1, 1, 1)
            else:
                payment_date = timezone.datetime(year, month + 1, 1)
        else:
            # Default to the current month’s next month if df3 is not set
            payment_date = timezone.now().replace(day=1) + timedelta(days=31)
            payment_date = payment_date.replace(day=1)

        # Define the parameters for the payment record
        payment_data = {
            'sales_challan_number': order.scno,
            'oid': order.oid,
            'product': order.product,
            'hsn': order.product.hsn_code,
            'amount': order.amount,
            'date': payment_date,
            'status': 'unpaid',
            'df2': order.df2
        }

        # Update or create the Payment record based on sales_challan_number
        payment, created = Payment.objects.update_or_create(
            sales_challan_number=order.scno,
            defaults=payment_data
        )

        # Redirect to the Payment page
        return redirect('payment_page_url')
    else:
        # If it's not a POST request, redirect to the show view
        return redirect('show_view_url')


@require_POST
def mark_payment_paid(request, payment_id):
    # Get the payment object based on payment_id
    payment = get_object_or_404(Payment, id=payment_id)

    # Update the status to 'paid'
    payment.status = 'paid'
    payment.save()

    # Redirect back to the payment page
    return redirect('payment_page_url')

def paymentView(request):
    # Get the status filter from the GET request parameters
    status_filter = request.GET.get('status', '')

    # Filter payments based on the status, if a filter is applied
    if status_filter:
        payments = Payment.objects.filter(status=status_filter)
    else:
        #payments = Payment.objects.all()  # Retrieve all records if no filter is specified
        payments = Payment.objects.annotate(
            unpaid_first=Case(
                When(status='unpaid', then=0),  # Give 'unpaid' a priority of 0
                default=1,  # All other statuses have a priority of 1
                output_field=IntegerField(),
            )
        ).order_by('unpaid_first', 'status')

    context = {
        'payments': payments,
    }
    return render(request, 'invoice/payment.html', context)

def delete_payment(request, payment_id):
    # Check if it's a POST request to ensure form submission
    if request.method == "POST":
        # Get the payment object and delete it
        payment = get_object_or_404(Payment, id=payment_id)
        payment.delete()
    # Redirect to the payment list page
    return redirect('payment_page_url')

def update_payment_date(request, payment_id):
    if request.method == "POST":
        # Get the payment object
        payment = get_object_or_404(Payment, id=payment_id)

        # Calculate the next month based on the current payment date
        current_date = payment.date
        year = current_date.year
        month = current_date.month

        # If it's December, move to January of the next year
        if month == 12:
            new_date = current_date.replace(year=year + 1, month=1, day=1)
        else:
            # Move to the first day of the next month
            new_date = current_date.replace(month=month + 1, day=1)

        # Update the payment date
        payment.date = new_date
        payment.save()

    # Redirect to the payment list page
    return redirect('payment_page_url')
def updateView(request, f_oid):
    obj = Orders.objects.get(oid=f_oid)
    form = OrderForm(instance=obj)
    if request.method == 'POST':
        form = OrderForm(request.POST, instance=obj)
        if form.is_valid():
            order_instance = form.save(commit=False)

            order_instance.amount = calculate_amount(order_instance)

            order_instance.save()

            return redirect('show_url')
            return redirect('show_url')
    template_name = 'invoice/order.html'
    context = {'form': form}
    return render(request, template_name, context)

def updateCusView(request, f_Cname):
    obj = Customer.objects.get(Cname=f_Cname)
    form = CustomerForm(instance=obj)
    if request.method == 'POST':
        form = CustomerForm(request.POST, instance=obj)
        if form.is_valid():
            form.save()
            return redirect('customer_info_url')
    template_name = 'invoice/customer_add.html'
    context = {'form': form}

    return render(request, template_name, context)

def deleteView(request, f_oid):
    obj = Orders.objects.get(oid=f_oid)
    if request.method == 'POST':
        obj.delete()
        return redirect('show_url')
    template_name = 'invoice/confirmation.html'
    context = {'obj': obj}
    return render(request, template_name, context)

def deleteCusView(request, f_Cname):
    obj = Customer.objects.get(Cname=f_Cname)
    if request.method == 'POST':
        obj.delete()
        return redirect('customer_info_url')
    template_name = 'invoice/CusCustomer.html'
    context = {'obj': obj}
    return render(request, template_name, context)

def export_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="orders.csv"'

    writer = csv.writer(response)
    writer.writerow(['Order ID', 'Supplier Name', 'Place', 'Driver Name', 'Lorry number', 'Purchase Challan no.', 'Sales Challan no.', 'Product Name', 'pcs/Fts', 'Number of trips', 'Date Field 1', 'Date Field 2', 'Date Field 3'])

    orders = Orders.objects.all()
    for order in orders:
        writer.writerow([order.oid, order.fname, order.Pname, order.Dname, order.Lno, order.pcno, order.scno, order.Product, order.pcs, order.trip, order.df1, order.df2, order.df3])

    return response


def customer_add(request):
    if request.method == 'POST':
        form = CustomerForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('customer_info_url')  # Redirect to the customer info page
    else:
        form = CustomerForm()

    return render(request, 'invoice/customer_add.html', {'form': form})


def customer_info(request):
    # Get the search query from the URL parameter 'q'
    query = request.GET.get('q', '')

    # If a query is provided, filter customers based on the name
    if query:
        customers = Customer.objects.filter(Q(Cname__icontains=query))
    else:
        # If no query, fetch all customers
        customers = Customer.objects.all()

    # Render the customerInfo.html template with the customers and query
    template_name = 'invoice/customerInfo.html'
    context = {'customers': customers, 'query': query}
    return render(request, template_name, context)


def supplier_add(request):
    if request.method == 'POST':
        form = SupplierForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('supplier_info_url')  # Redirect to the supplier info page
    else:
        form = SupplierForm()

    return render(request, 'invoice/supplier_add.html', {'form': form})


# Display supplier information
def supplier_info(request):
    query = request.GET.get('q', '')  # Get the search query from the URL parameter 'q'

    if query:
        suppliers = Supplier.objects.filter(Q(fname__icontains=query))
    else:
        suppliers = Supplier.objects.all()

    template_name = 'invoice/supplierInfo.html'
    context = {'suppliers': suppliers, 'query': query}
    return render(request, template_name, context)


# Update existing supplier
def updateSupplierView(request, fname):
    supplier = get_object_or_404(Supplier, fname=fname)
    form = SupplierForm(instance=supplier)
    if request.method == 'POST':
        form = SupplierForm(request.POST, instance=supplier)
        if form.is_valid():
            form.save()
            return redirect('supplier_info_url')
    template_name = 'invoice/supplier_add.html'
    context = {'form': form}
    return render(request, template_name, context)


# Delete supplier
def deleteSupplierView(request, fname):
    supplier = get_object_or_404(Supplier, fname=fname)
    if request.method == 'POST':
        supplier.delete()
        return redirect('supplier_info_url')
    template_name = 'invoice/SupplierConfirmation.html'
    context = {'supplier': supplier}
    return render(request, template_name, context)

def lorry_add(request):
    if request.method == 'POST':
        form = LorryForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('lorry_info_url')  # Redirect to the lorry info page
    else:
        form = LorryForm()

    return render(request, 'invoice/lorry_add.html', {'form': form})


# Display lorry information
def lorry_info(request):
    query = request.GET.get('q', '')  # Get the search query from the URL parameter 'q'

    if query:
        lorries = Lorry.objects.filter(Q(Lno__icontains=query))  # Filter by lorry number
    else:
        lorries = Lorry.objects.all()

    template_name = 'invoice/lorryInfo.html'
    context = {'lorries': lorries, 'query': query}
    return render(request, template_name, context)


# Update existing Lorry
def updateLorryView(request, Lno):
    lorry = get_object_or_404(Lorry, Lno=Lno)
    form = LorryForm(instance=lorry)
    if request.method == 'POST':
        form = LorryForm(request.POST, instance=lorry)
        if form.is_valid():
            form.save()
            return redirect('lorry_info_url')
    template_name = 'invoice/lorry_add.html'
    context = {'form': form}
    return render(request, template_name, context)


# Delete Lorry
def deleteLorryView(request, Lno):
    lorry = get_object_or_404(Lorry, Lno=Lno)
    if request.method == 'POST':
        lorry.delete()
        return redirect('lorry_info_url')
    template_name = 'invoice/LorryConfirmation.html'
    context = {'lorry': lorry}
    return render(request, template_name, context)

def product_add(request):
    if request.method == 'POST':
        form = ProductForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('product_info_url')  # Redirect to the product info page
    else:
        form = ProductForm()

    return render(request, 'invoice/product_add.html', {'form': form})


# Display product information
def product_info(request):
    query = request.GET.get('q', '')  # Get the search query from the URL parameter 'q'
    hsn_query = request.GET.get('hsn', '')
    filters = Q()
    if query:
        filters &= Q(name__icontains=query)
    if hsn_query:
        filters &= Q(hsn_code__icontains=hsn_query)

    products = Product.objects.filter(filters)

    template_name = 'invoice/productInfo.html'
    context = {'products': products, 'query': query}
    return render(request, template_name, context)


# Update existing product
def updateProductView(request, name):
    product = get_object_or_404(Product, name=name)
    form = ProductForm(instance=product)
    if request.method == 'POST':
        form = ProductForm(request.POST, instance=product)
        if form.is_valid():
            form.save()
            return redirect('product_info_url')
    template_name = 'invoice/product_add.html'
    context = {'form': form}
    return render(request, template_name, context)


# Delete product
def deleteProductView(request, name):
    product = get_object_or_404(Product, name=name)
    if request.method == 'POST':
        product.delete()
        return redirect('product_info_url')
    template_name = 'invoice/ProductConfirmation.html'
    context = {'product': product}
    return render(request, template_name, context)


def get_product_details(request, product_id):
    try:
        # Query the Product model to get details based on the product_id
        product = Product.objects.get(name=product_id)

        # Example response data
        data = {
            'name': product.name,
            'rate': product.rate,
            'quantity_per': product.quantity_per,
            # Add other fields as needed
        }

        return JsonResponse(data)
    except Product.DoesNotExist:
        return JsonResponse({'error': 'Product not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def generate_invoice_id(model_name,invoice_end_date):
    current_month = invoice_end_date.strftime("%b").upper()
    current_monthInt = invoice_end_date.month
    current_year = invoice_end_date.year

    # Determine financial year
    financial_year = (
        f"{current_year - 1}-{current_year % 100}"
        if current_monthInt < 4
        else f"{current_year}-{(current_year + 1) % 100}"
    )
    distinct_records = model_name.objects.filter(Invoice__endswith=f"/{financial_year}").values("Invoice").distinct()

# Display the distinct Invoice values
    print("Distinct Invoice Numbers:")
    for record in distinct_records:
        print(record["Invoice"])

    # Count the distinct records
    invoice_count = len(distinct_records)
    print("Calculated Invoice Count:", invoice_count)

    # Generate the invoice ID
    invoice_id = f"GT/{invoice_count + 1}/{current_month}/{financial_year}"

    return invoice_id


def billView(request):
    # Initialize bill_items before handling the request
    print("Request Method:", request.method)

    logger = logging.getLogger(__name__)
    bill_items = []
    start_date = None
    end_date = None
    orders = Orders.objects.none()
    #Bill Submit button
    if request.method == "POST" and request.POST.get('action') == 'bill':
        selected_temp_rate_ids = request.POST.getlist('selected_rows')
        print("Selected Temp Rate IDs:", selected_temp_rate_ids)


        if selected_temp_rate_ids:
            # Fetch TempRate objects based on selected IDs
            selected_temp_rates = TempRate.objects.filter(id__in=selected_temp_rate_ids)
            print("Selected Temp Rates:", selected_temp_rates)

            # Merge their csoid values into a list
            merged_csoid_list = []
            for temp_rate in selected_temp_rates:
                merged_csoid_list.extend(temp_rate.csoid.split(','))

            # Process the merged OID list for billing or further actions
            print("Merged OID list:", merged_csoid_list)
            Orders.objects.filter(oid__in=merged_csoid_list).update(billed='Yes')

            temp_table = TempTable.objects.get(pk=1)
            invoice_number = generate_invoice_id(Record, temp_table.end_date)

            for oid in merged_csoid_list:
                # Fetch data from Orders table
                order_data = Orders.objects.get(oid=oid)

                # Fetch data from Product table
                product_data = Product.objects.get(name=order_data.product.name)

                # Fetch data from Billing table
                billing_data = Billing.objects.get(oid=oid)

                customer_data = Customer.objects.get(Cname=order_data.Cname)

                temp_table = TempTable.objects.get(pk=1)
                end_date_from_temp_table = temp_table.end_date
                # Additional data for the Record table
                Cname = customer_data
                Cadr = customer_data.adr
                Sadr = order_data.Pname
                state = customer_data.state
                code = customer_data.code
                gst = customer_data.gst
                date = order_data.df3
                lorry_no = order_data.Lno
                trip = order_data.trip
                challan = order_data.scno
                hsn = product_data.hsn_code
                tax_rate = product_data.tax_rate
                product = order_data.product.name
                pcs = order_data.pcs
                rate = billing_data.final_rate
                amount = order_data.amount
                cgst = amount * (product_data.tax_rate / 200)
                sgst = amount * (product_data.tax_rate / 200)
                final_amount = amount + cgst + sgst
                purchaseRate = order_data.purchaseRate

                print(final_amount)
                # Create a Record object and save it to the database
                record = Record.objects.create(
                    Cname = Cname,
                    Cadr = Cadr,
                    Sadr = Sadr,
                    state = state,
                    code = code,
                    gst = gst,
                    date=date,
                    InvoiceEnd=end_date_from_temp_table,
                    Invoice = invoice_number,
                    lorry_no=lorry_no,
                    trip=trip,
                    challan=challan,
                    hsn=hsn,
                    tax_rate = tax_rate,
                    product=product,
                    pcs=pcs,
                    rate=rate,
                    amount=amount,
                    cgst=cgst,
                    sgst=sgst,
                    final_amount=final_amount,
                    purchaseRate = purchaseRate,
                )
            filtered_records = Record.objects.filter(Invoice=invoice_number)

# Calculate the sum of the final amount for the filtered records
            total_amount_sum = filtered_records.aggregate(Sum('final_amount'))['final_amount__sum']

# Update the 'total_amount' column for the filtered records
            filtered_records.update(total_amount=total_amount_sum)

            return redirect('showBill_url')
    #if request.method == 'POST':
        # Clear the TempRate table on every POST request
    TempRate.objects.all().delete()
    selected_customer = request.GET.get('customer', '')


    # Handle the case where a date range is selected
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    # Parse dates if provided
    try:
        start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
    except (TypeError, ValueError):
        start_date = None

    try:
        end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
    except (TypeError, ValueError):
        end_date = None

    temp_table, created = TempTable.objects.get_or_create(pk=1)
    temp_table.end_date = end_date
    temp_table.save()
    orders = Orders.objects.filter(billed='No')
    if start_date and end_date:
        # Filter orders within the date range
        #orders = Orders.objects.filter(df3__range=(start_date, end_date), billed='No')
        orders = orders.filter(df3__range=(start_date, end_date))
    if selected_customer:
        # Filter by customer name (adjust field name to match your Orders model)
        #orders = orders.filter(Cname=selected_customer)
        orders = orders.filter(Q(Cname__Cname__icontains=selected_customer))
    if not orders.exists():
        orders = Orders.objects.filter(billed='No')

    # Group and aggregate by customer, location, and product


    valid_orders = orders.filter(df3__isnull=False).order_by('Cname', 'Pname', 'product', 'df3')

# Group and aggregate by customer, location, product, month, and year
    for key, group in groupby(
        valid_orders.values('Cname', 'Pname', 'product', 'trip', 'pcs', 'amount', 'oid', 'df3'),
        key=lambda x: (x['Cname'], x['Pname'], x['product'], x['df3'].month, x['df3'].year)  # Group by month and year
    ):
        group_list = list(group)
        if group_list:
            merged_oids = ','.join(str(item['oid']) for item in group_list)
            total_trips = sum(item['trip'] for item in group_list)
            total_quantity = sum(item['pcs'] for item in group_list)
            total_amount = sum(item['amount'] for item in group_list)

            oids_list = merged_oids.split(',')
            are_all_rated = all(Billing.objects.filter(oid=oid).exists() for oid in oids_list)

            rated_value = 'Yes' if are_all_rated else 'No'

            # Insert aggregated data into TempRate
            temp_rate = TempRate.objects.create(
                customer_name=key[0],  # Customer name
                location=key[1],       # Location
                product=key[2],        # Product
                csoid=merged_oids      # Merged OIDs
            )

            # Include TempRate id in bill_item for linking purposes
            bill_items.append({
                'Cname': key[0],
                'Pname': key[1],
                'product': key[2],
                'csoid': merged_oids,
                'total_trips': total_trips,
                'total_pcs': total_quantity,
                'total_amount': total_amount,
                'temp_rate_id': temp_rate.id,  # Include TempRate ID
                'month': key[3],  # Month extracted from df3
                'year': key[4],    # Year extracted from df3
                'merged_count': len(merged_oids.split(',')),
                'rated': rated_value
            })



    # Redirect or render the page after processing
    logger.debug("Bill Items: %s", bill_items)
    print(bill_items)
        #return redirect('show_url')

    # Handle GET requests and populate the form
    template_name = 'invoice/Bill.html'
    selected_customer = request.GET.get('customer', '')  # Moved this line here

    context = {
        'customer_names': Customer.objects.values_list('Cname', flat=True).distinct(),
        'bill_items': bill_items,  # The aggregated bill items you want to display
        'selected_customer': selected_customer,  # Customer selected by the user
        'start_date': start_date,  # Start date for filtering
        'end_date': end_date,  # End date for filtering
    }

    return render(request, template_name, context)




def rateView(request, id):
    # Fetch the TempRate object using the passed 'id' (which is the TempRate id)
    try:
        temp_rate = TempRate.objects.get(id=id)
    except TempRate.DoesNotExist:
        # Handle the case where TempRate does not exist
        return redirect('error_url')  # Redirect to an error page or handle gracefully

    # Get the comma-separated OIDs from the TempRate object
    csoid = temp_rate.csoid
    oid_list = csoid.split(',')  # Split the csoid into a list of OIDs

    # Fetch the orders corresponding to the OIDs in csoid
    orders = Orders.objects.filter(oid__in=oid_list)
    aggregated_quantity = orders.aggregate(Sum('pcs'))['pcs__sum']
    aggregated_amount = orders.aggregate(Sum('amount'))['amount__sum']

    if orders.exists():
        first_order = orders.first()
        customer_name = first_order.Cname
        bill_date = first_order.df1


    if request.method == "POST":
        print("POST request received")
        total_rate = float(request.POST.get('total_rate'))  # Get the total rate from the form
        oid = request.POST.get('oid')
        product = request.POST.get('product')
        place = request.POST.get('place')
        material_rate = request.POST.get('material_rate')
        transport_rate = total_rate - float(material_rate)
        # Update all orders with the new rate * quantity
        for order in orders:
            product = order.product
            amount = order.pcs * total_rate / product.quantity_per
            order.amount = amount  # Update the aggregated amount
            order.save()  # Save changes to the database

        # Display customer name and date of the first order
        if orders.exists():
            first_order = orders.first()
            customer_name = first_order.Cname  # Assuming Cname is the customer name field
            bill_date = first_order.df1  # Assuming df1 is the bill date field
            customer_instance = Customer.objects.get(Cname=customer_name)

        if 'action' in request.POST and request.POST['action'] == 'submit':


            for order in orders:
                existing_billing = Billing.objects.filter(oid=order.oid)
                print(f"Processing Order ID: {order.oid}, Existing Billing Count: {existing_billing.count()}")

                if existing_billing.exists():
            # Delete the existing billing records
                    existing_billing.delete()
                    print(f"Deleted existing billing records for Order ID: {order.oid}")

                billing_instance = Billing(
                Cname=customer_instance,
                bill_date= bill_date,
                final_rate=total_rate,
                material_rate=material_rate,
                transport_rate=transport_rate,
                final_amount=order.amount,
                oid=order.oid,
                product=order.product,
                place= order.Pname
                )

            #print("Billing Data:", billing_instance.product)
                billing_instance.save()
            #print("Billing Data:", billing_instance.oid)
                print(f"Billing Data Saved: OID: {billing_instance.oid}, Product: {billing_instance.product}")

        # Fetch previous bills based on customer name and product


        # Pass the necessary context data to the template
        context = {
            'customer_name': customer_name,
            'bill_date': bill_date,
            'orders': orders,

        }
        try:
            return render(request, 'invoice/bill.html', context)
        except TemplateDoesNotExist:
            return redirect('bill_url')
        #return render(request, 'invoice/bill.html', context)
    previous_bills = Billing.objects.filter(
            Cname=first_order.Cname,
            product=first_order.product
        ).order_by('-bill_date')
    # If the request method is not POST, just display the rate page with current data
    context = {
        'orders': orders,
        'customer_name': temp_rate.customer_name,  # You can use fields from TempRate here
        'location': temp_rate.location,
        'product': temp_rate.product,
        'aggregated_quantity': aggregated_quantity,
        'aggregated_amount': aggregated_amount,
        'temp_rate': temp_rate,
        'previous_bills': previous_bills,
        'quantity_per': first_order.product.quantity_per
    }
    print("Previous Bill",previous_bills)
    print(orders)
    return render(request, 'invoice/rate.html', context)

def showBillView(request):
    customer_query = request.GET.get('customer', '')
    date_query = request.GET.get('date', '')

    # Retrieve distinct records based on the 'Invoice' field
    records = Record.objects.values('Invoice').annotate(
        record_id=Max('record_id'),
        date=Max('date'),
        customer=Max('Cname'),
        total_amount=Round(Max('total_amount')),
    ).order_by('-date')

    if customer_query:
        records = records.filter(Q(customer__icontains=customer_query))
    if date_query:
        date = datetime.strptime(date_query, '%Y-%m-%d').date()
        records = records.filter(date=date)

    template_name = 'invoice/showBill.html'
    context = {'records': records}
    return render(request, template_name, context)



def replace_placeholders(template_path, output_path, replacements):
    template = MailMerge(template_path)
    template.merge(**replacements)
    template.write(output_path)

def merge_all_cells_in_row(row):
    """
    Merges all the cells in the given row, making the row appear without table lines.
    """
    # Merge all cells in the row
    first_cell = row.cells[0]
    for cell in row.cells[1:]:
        first_cell.merge(cell)

def set_borders(cell, top=None, bottom=None, left="single", right="single"):
    """Set specific borders for a cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Define and set the individual borders
    for border_name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), value if value else 'nil')
        tcBorders.append(border)

def add_rows_to_table(doc, query_set):
    # Access the predefined table (assuming it's the first table in the document)
    table = doc.tables[0]
    total_cgst = 0
    total_sgst = 0
    total_trip = 0
    pretax_amount = 0
    tax_rate = 0

    # Initialize variables to track cumulative data for similar products
    cumulative_values = {'pcs': 0, 'amount': 0, 'cgst': 0, 'sgst': 0, 'final_amount': 0}
    previous_product = None  # Track the previous product to detect duplicates
    product_count = 0
    product_rate = 0
    duplicate_row_indices = defaultdict(list)


    date_column_index = 0  # Assuming date is the first column
    for row in table.rows:
        row.cells[date_column_index].width = Inches(0.6)

    # Iterate through the queryset and add rows to the table
    for index, record in enumerate(query_set):

        # Add a new row to the table
        if record.product == previous_product:
        # Increment count and add to cumulative values
            product_count += 1
            cumulative_values['pcs'] += record.pcs
            cumulative_values['amount'] += record.amount
            cumulative_values['cgst'] += record.cgst
            cumulative_values['sgst'] += record.sgst
            cumulative_values['final_amount'] += record.final_amount
            duplicate_row_indices[record.product].append(len(table.rows) - 1)
        else:
        # If the previous product appeared more than once, add a cumulative row
            if product_count > 1:
                cumulative_row_cells = table.add_row().cells
                cumulative_row_cells[6].text = str(int(cumulative_values['pcs']))
                cumulative_row_cells[7].text = str(int(product_rate))
                cumulative_row_cells[8].text = str(cumulative_values['amount'])
                cumulative_row_cells[9].text = f"{round(cumulative_values['cgst'], 2)} ({str(tax_rate / 2)}%)"
                cumulative_row_cells[10].text = f"{round(cumulative_values['sgst'], 2)} ({str(tax_rate / 2)}%)"
                cumulative_row_cells[11].text = str(cumulative_values['final_amount'])
                for cell in cumulative_row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    set_borders(cell, top="nil", bottom="nil")


            empty_row = table.add_row()
            for cell in empty_row.cells:
                set_borders(cell, top="nil", bottom="nil")

            # Reset cumulative values and product count for the new product
            previous_product = record.product
            product_count = 1  # Reset count to 1 for the new product
            product_rate = record.rate
            tax_rate = record.tax_rate
            cumulative_values = {
                'pcs': record.pcs,
                'amount': record.amount,
                'cgst': record.cgst,
                'sgst': record.sgst,
                'final_amount': record.final_amount
            }


        row_cells = table.add_row().cells

        row_idx = len(table.rows) - 1  # Track the new row index
        duplicate_row_indices[record.product].append(row_idx)  # Track duplicates by product

        # Populate the cells with data from the queryset
        #row_cells[0].text = record.date.strftime("%b %d") #str(record.date)
        row_cells[0].text = record.date.strftime("%d/%m")
        row_cells[1].text = record.lorry_no
        row_cells[2].text = str(int(record.trip))
        row_cells[3].text = record.challan
        row_cells[4].text = record.hsn
        row_cells[5].text = record.product
        row_cells[6].text = str(int(record.pcs))

        row_cells[7].text = str(int(record.rate))
        row_cells[8].text = str(record.amount)
        row_cells[9].text = f"{round(record.cgst, 2)} ({str(record.tax_rate / 2)}%)"
        row_cells[10].text = f"{round(record.sgst, 2)} ({str(record.tax_rate / 2)}%)"
        row_cells[11].text = str(round(record.final_amount,2))
        #final_amount_text = str(round(record.final_amount))  # Round the final amount

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_borders(cell, top="nil", bottom="nil")


    if product_count > 1:
        cumulative_row_cells = table.add_row().cells
        cumulative_row_cells[6].text = str(int(cumulative_values['pcs']))
        cumulative_row_cells[7].text = str(int(product_rate))
        cumulative_row_cells[8].text = str(cumulative_values['amount'])
        cumulative_row_cells[9].text = f"{round(cumulative_values['cgst'], 2)} ({str(tax_rate / 2)}%)"
        cumulative_row_cells[10].text = f"{round(cumulative_values['sgst'], 2)} ({str(tax_rate / 2)}%)"
        cumulative_row_cells[11].text = str(cumulative_values['final_amount'])
        for cell in cumulative_row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_borders(cell, top="nil", bottom="nil")

        row_idx = len(table.rows) - 1  # This is the index of the cumulative row
        duplicate_row_indices[previous_product].append(row_idx)
    for product, indices in duplicate_row_indices.items():
        # Skip clearing if there was only one row for this product
        if len(indices) > 1:
            for row_idx in indices:
                for col_idx in range(7, 12):  # Adjust columns based on table structure
                    table.rows[row_idx].cells[col_idx].text = ""



    for index, record in enumerate(query_set):
        total_cgst += record.cgst
        total_trip += record.trip
        total_sgst += record.sgst
        pretax_amount += record.amount

    # If it's the last record, add extra empty rows and total calculations
        if index == len(query_set) - 1:
            for _ in range(2):
                extra_empty_row = table.add_row()
                for cell in extra_empty_row.cells:
                    set_borders(cell, top="nil", bottom="nil")

            empty_row_1 = table.add_row()
            for cell in empty_row_1.cells:
                set_borders(cell, top="nil")

            total_amount_row = table.add_row().cells

            total_trip_no_cell = total_amount_row[2]
            total_trip_no_cell.text = str(int(total_trip))

            total_label_cell = total_amount_row[6].merge(total_amount_row[7])
            total_label_cell.text = 'Total Rs'
            total_label_run = total_label_cell.paragraphs[0].runs[0]
            total_label_run.bold = True
            total_label_run.font.size = Pt(13)
            for paragraph in total_label_cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            pretax_amount_cell = total_amount_row[8]
            pretax_amount_cell.text = f"{round(pretax_amount, 2)}"
            for paragraph in pretax_amount_cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            total_cgst_cell = total_amount_row[9]
            total_cgst_cell.text = str(round(total_cgst, 2))
            for paragraph in total_cgst_cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            total_sgst_cell = total_amount_row[10]
            total_sgst_cell.text = str(round(total_sgst, 2))
            for paragraph in total_sgst_cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            total_amount_cell = total_amount_row[11]
            total_amount_cell.text = str(int(record.total_amount))
            for paragraph in total_amount_cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            total_amount_run = total_amount_cell.paragraphs[0].runs[0]
            total_amount_run.bold = True
            total_amount_run.underline = True
            total_amount_run.font.size = Pt(12.5)

            # Add a new row for 'Rupees' and the total amount in words
            rupees_row = table.add_row().cells

            # Add the total amount in words
            words = num2words(round(record.total_amount), lang='en_IN')
            words_no_commas = words.replace(',', '').replace(' and ', ' ').replace('-', ' ')
            capitalized_words = 'Rupees :' + words_no_commas.title() + ' Only.'
            amount_in_words_cell = rupees_row[0].merge(rupees_row[11])
            amount_in_words_cell.text = capitalized_words

def printDocxView(request):
    if request.method == 'POST':
        invoice_id = request.POST.get('invoice_id')

        # Load the DOCX template file
#        template_path = 'invoiceApp/templates/invoice/InVoice.docx'
        template_path = os.path.join(settings.BASE_DIR, 'invoiceApp/templates/invoice/InVoice.docx')
        output_path = f'modified_invoice_{invoice_id.replace("/", "_")}.docx'

        # Get records based on the invoice_id
        record = Record.objects.filter(Invoice=invoice_id).first()

        # Convert all fields to strings
        Cname = str(record.Cname)
        state = str(record.state)
        code = str(int(record.code))
        Cadr = str(record.Cadr)
        Sadr = str(record.Sadr)
        gst = str(record.gst)
        date = str(record.InvoiceEnd.strftime("%d/%m/%Y"))


        # Define replacements for placeholders
        replacements = {
            'Cname': Cname,
            'invoice_id': invoice_id,
            'Cadr': Cadr,
            'state': state,
            'code': code,
            'gst': gst,
            'Sadr': Sadr,
            'date': date,

        }
        replace_placeholders(template_path, output_path, replacements)

        doc = Document(output_path)

        # Your queryset
        query_set = Record.objects.filter(Invoice=invoice_id)

        # Add rows to the table
        add_rows_to_table(doc, query_set )

        # Save the modified DOCX document
        doc.save(output_path)

        # Save the modified DOCX document to a BytesIO buffer
        buffer = io.BytesIO()
        with open(output_path, 'rb') as file:
            buffer.write(file.read())

        # Create a FileResponse and return it for download
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=modified_invoice_{invoice_id.replace("/", "_")}.docx'
        response.write(buffer.getvalue())
        return response

def analytics_dashboard(request):
    return render(request, 'invoice/analytics.html')  # The HTML file should be inside the templates folder

def total_revenue_by_product(request):
    # Aggregate total revenue by product
    revenue_data = (
        Orders.objects
        .values('product__name')
        .annotate(total_revenue=Sum('amount'))
        .order_by('product__name')
    )

    return JsonResponse(list(revenue_data), safe=False)

def average_rate_per_product(request):
    # Aggregate average final rate by product from the Billing table
    rate_data = (
        Billing.objects
        .values('product')
        .annotate(avg_rate=Avg('final_rate'))  # Taking the average of final_rate
        .order_by('product')
    )

    return JsonResponse(list(rate_data), safe=False)

def top_customer_contribution(request):
    # Aggregate total contribution by customer
    customer_data = (
        Orders.objects
        .values('Cname')
        .annotate(total_contribution=Sum('amount'))
        .order_by('-total_contribution')[:5]  # Top 5 customers
    )

    return JsonResponse(list(customer_data), safe=False)

def tax_statement_view(request):
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    hsn_code_filter = request.GET.get('hsn_code')

    # If dates are provided, parse them, else set to None
    if start_date:
        start_date = parse_date(start_date)
    if end_date:
        end_date = parse_date(end_date)

    # Initialize an empty dictionary to hold tax statements per product
    tax_statements = {}



    # Query all records
    records = Record.objects.all()
    #records = Record.objects.exclude(challan__in = unpaid_challans)
    if start_date and end_date:
        records = records.filter(date__range=[start_date, end_date])
    if hsn_code_filter:
        matching_products = Product.objects.filter(hsn_code=hsn_code_filter).values_list('name', flat=True)
        # Now filter records based on the product names
        records = records.filter(product__in=matching_products)

    total_taxable_amount = 0
    total_cgst = 0
    total_sgst = 0
    total_final_amount = 0
    total_net_profit = 0


    # Loop through each record and aggregate the data
    for record in records:
        product = record.product

        try:
            product_obj = Product.objects.get(name=product)
            hsn_code = product_obj.hsn_code
            quantity_per = product_obj.quantity_per
        except Product.DoesNotExist:
            hsn_code = None
            quantity_per = 1

        # Initialize if product not already in the dictionary
        if product not in tax_statements:
            tax_statements[product] = {
                'hsn_code': hsn_code,
                'total_taxable_amount': 0,
                'total_cgst': 0,
                'total_sgst': 0,
                'total_final_amount': 0,
                'total_quantity': 0,
                'total_net_profit': 0,
            }

        if quantity_per > 0:  # Avoid division by zero
            purchaseAmount = (((record.purchaseRate or 0) * record.pcs)/quantity_per) *(1+ (record.tax_rate/100))
            net_profit = record.final_amount - purchaseAmount
            #net_profit = ((record.rate - (record.purchaseRate or 0)) * record.pcs) / quantity_per
        else:
            net_profit = 0

        # Aggregate the amounts
        tax_statements[product]['total_taxable_amount'] += record.amount
        tax_statements[product]['total_cgst'] += record.cgst
        tax_statements[product]['total_sgst'] += record.sgst
        tax_statements[product]['total_final_amount'] += record.final_amount
        tax_statements[product]['total_quantity'] += record.pcs
        tax_statements[product]['total_net_profit'] += net_profit

        total_taxable_amount += record.amount
        total_cgst += record.cgst
        total_sgst += record.sgst
        total_final_amount += record.final_amount
        total_net_profit += net_profit


    tax_statements_by_customer = {}


    # Loop through each record and aggregate the data by customer and product
    for record in records:
        customer = record.Cname  # Foreign key to Customer
        product = record.product

        try:
            product_obj = Product.objects.get(name=product)
            hsn_code = product_obj.hsn_code
        except Product.DoesNotExist:
            hsn_code = None

        # Initialize if customer not already in the dictionary
        if customer not in tax_statements_by_customer:
            tax_statements_by_customer[customer] = {
                'gst': customer.gst,
                'products': {}
            }

        # Initialize product if not already in the customer's dictionary
        if product not in tax_statements_by_customer[customer]['products']:
            tax_statements_by_customer[customer]['products'][product] = {
                'hsn_code': hsn_code,
                'total_taxable_amount': 0,
                'total_cgst': 0,
                'total_sgst': 0,
                'total_final_amount': 0,
                'total_quantity': 0,
            }

        # Aggregate the amounts for this customer and product
        tax_statements_by_customer[customer]['products'][product]['total_taxable_amount'] += record.amount
        tax_statements_by_customer[customer]['products'][product]['total_cgst'] += record.cgst
        tax_statements_by_customer[customer]['products'][product]['total_sgst'] += record.sgst
        tax_statements_by_customer[customer]['products'][product]['total_final_amount'] += record.final_amount
        tax_statements_by_customer[customer]['products'][product]['total_quantity'] += record.pcs

    for customer, products in tax_statements_by_customer.items():
        for product, values in products['products'].items():
            values['total_quantity'] = round(values['total_quantity'])


    for product, values in tax_statements.items():
        values['total_quantity'] = round(values['total_quantity'])

    purchase_statements = {}
    unpaid_oids = Payment.objects.values_list('oid', flat=True)

    # Get purchase records from the Orders table excluding the OIDs present in the Payment table
    order_records = Orders.objects.exclude(oid__in=unpaid_oids)

    # Get purchase records from the Orders table filtered by the unpaid OIDs
    if start_date and end_date:
        order_records = order_records.filter(df2__range=[start_date, end_date])

    paid_order_records = Orders.objects.filter(
        Q(oid__in=Payment.objects.filter(status='paid', date__range=[start_date, end_date]).values_list('oid', flat=True)))

    # Combine the two querysets
    combined_order_records = order_records | paid_order_records

    # Aggregated totals for purchase statement
    total_purchase_pretax_amount = 0
    total_purchase_cgst = 0
    total_purchase_sgst = 0
    total_purchase_amount_Overall = 0


    for order in combined_order_records:
        product = order.product.name

        try:
            product_obj = Product.objects.get(name=product)
            quantity_per = product_obj.quantity_per
            hsn_code = product_obj.hsn_code
        except Product.DoesNotExist:
            quantity_per = 1
            hsn_code = None

        # Use the purchaseRate from the order
        purchase_rate = order.purchaseRate

        if purchase_rate is None or quantity_per is None:
            print(f"Order ID: {order.oid} has None values: purchase_rate={purchase_rate}, quantity_per={quantity_per}")
            continue


        purchase_pretax_amount = purchase_rate * (order.pcs / quantity_per)  # Calculate purchase pretax amount
        purchase_cgst = purchase_pretax_amount * product_obj.tax_rate / 200
        purchase_sgst = purchase_pretax_amount * product_obj.tax_rate / 200
        total_purchase_amount = purchase_pretax_amount + purchase_cgst + purchase_sgst

        if product not in purchase_statements:
            purchase_statements[product] = {
                'hsn_code': hsn_code,
                'total_purchase_amount': 0,
                'total_purchase_pretax_amount': 0,
                'total_purchase_cgst': 0,
                'total_purchase_sgst': 0,
                'total_quantity': 0,
            }

        purchase_statements[product]['total_purchase_amount'] += total_purchase_amount
        purchase_statements[product]['total_purchase_pretax_amount'] += purchase_pretax_amount
        purchase_statements[product]['total_purchase_cgst'] += purchase_cgst
        purchase_statements[product]['total_purchase_sgst'] += purchase_sgst
        purchase_statements[product]['total_quantity'] += order.pcs

         # Update aggregated totals for overall purchase statement
        total_purchase_pretax_amount += purchase_pretax_amount
        total_purchase_cgst += purchase_cgst
        total_purchase_sgst += purchase_sgst
        total_purchase_amount_Overall += total_purchase_amount


    sorted_tax_statements = sorted(
        tax_statements.items(),
        key=lambda item: (item[1]['hsn_code'], item[0])  # Sort by HSN code, then by product name
    )

    # Convert the sorted list back to a dictionary if needed, or keep it as a list of tuples
    sorted_tax_statements_dict = {product: data for product, data in sorted_tax_statements}

    sorted_purchase_statements = sorted(
        purchase_statements.items(),
        key=lambda item: (item[1]['hsn_code'], item[0])
    )

    # Convert the sorted list back to a dictionary if needed, or keep it as a list of tuples
    sorted_purchase_statements_dict = {product: data for product, data in sorted_purchase_statements}

    # Prepare the context for rendering
    context = {
        'tax_statements': sorted_tax_statements_dict,
        'tax_statements_by_customer': tax_statements_by_customer,
        'purchase_statements': sorted_purchase_statements_dict,
        'start_date': start_date,
        'end_date': end_date,
        'total_taxable_amount': total_taxable_amount,
        'total_cgst': total_cgst,
        'total_sgst': total_sgst,
        'total_final_amount': total_final_amount,
        'total_purchase_pretax_amount': total_purchase_pretax_amount,
        'total_purchase_cgst': total_purchase_cgst,
        'total_purchase_sgst': total_purchase_sgst,
        'total_purchase_amount': total_purchase_amount,
        'total_net_profit': total_net_profit,
        'total_purchase_amount_Overall': total_purchase_amount_Overall,
    }

    # Render the template with the context
    return render(request, 'invoice/tax_statement.html', context)

